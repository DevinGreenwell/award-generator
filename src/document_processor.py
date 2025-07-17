"""
Document processor for extracting text from PDF and Word documents.
"""

import os
import logging
from typing import Optional, Tuple
import PyPDF2
from docx import Document
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

# Allowed file extensions and their MIME types
ALLOWED_EXTENSIONS = {
    'pdf': ['application/pdf'],
    'doc': ['application/msword'],
    'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_EXTRACTED_LENGTH = 10000  # Maximum characters to extract


class DocumentProcessor:
    """Handles document upload and text extraction."""
    
    def __init__(self, upload_folder: str = "uploads"):
        self.upload_folder = upload_folder
        os.makedirs(upload_folder, exist_ok=True)
    
    def allowed_file(self, filename: str) -> bool:
        """Check if file extension is allowed."""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    
    def validate_file(self, file) -> Tuple[bool, str]:
        """Validate uploaded file."""
        if not file or file.filename == '':
            return False, "No file selected"
        
        if not self.allowed_file(file.filename):
            return False, "Invalid file type. Only PDF and Word documents are allowed"
        
        # Check file size (Flask's request.files objects don't have a direct size property)
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)  # Reset file pointer
        
        if file_size > MAX_FILE_SIZE:
            return False, f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
        
        return True, "Valid"
    
    def extract_text_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                # Extract text from each page
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                
                # Clean up the text
                text = text.strip()
                
                # Limit text length
                if len(text) > MAX_EXTRACTED_LENGTH:
                    text = text[:MAX_EXTRACTED_LENGTH] + "... [Content truncated]"
                
                return text
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return None
    
    def extract_text_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Clean up the text
            text = text.strip()
            
            # Limit text length
            if len(text) > MAX_EXTRACTED_LENGTH:
                text = text[:MAX_EXTRACTED_LENGTH] + "... [Content truncated]"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return None
    
    def extract_text_from_doc(self, file_path: str) -> Optional[str]:
        """Extract text from DOC file (legacy format)."""
        # For now, we'll treat .doc files as .docx
        # Full .doc support would require python-docx2txt or similar
        return self.extract_text_from_docx(file_path)
    
    def process_file(self, file) -> Tuple[bool, str, Optional[str]]:
        """Process uploaded file and extract text."""
        # Validate file
        is_valid, message = self.validate_file(file)
        if not is_valid:
            return False, message, None
        
        # Secure filename
        filename = secure_filename(file.filename)
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        # Save file temporarily
        file_path = os.path.join(self.upload_folder, filename)
        try:
            file.save(file_path)
            
            # Extract text based on file type
            if file_extension == 'pdf':
                extracted_text = self.extract_text_from_pdf(file_path)
            elif file_extension == 'docx':
                extracted_text = self.extract_text_from_docx(file_path)
            elif file_extension == 'doc':
                extracted_text = self.extract_text_from_doc(file_path)
            else:
                return False, "Unsupported file type", None
            
            # Clean up - remove temporary file
            os.remove(file_path)
            
            if extracted_text:
                # Clean up whitespace
                extracted_text = ' '.join(extracted_text.split())
                
                # Create summary message
                word_count = len(extracted_text.split())
                message = f"Successfully extracted {word_count} words from {filename}"
                
                return True, message, extracted_text
            else:
                return False, "Could not extract text from file", None
                
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            # Clean up on error
            if os.path.exists(file_path):
                os.remove(file_path)
            return False, f"Error processing file: {str(e)}", None
    
    def analyze_document_for_achievements(self, text: str) -> str:
        """Analyze document text to extract achievement information."""
        try:
            from openai_client import OpenAIClient
            client = OpenAIClient()
            
            # Create a specific prompt for analyzing the document
            analysis_prompt = """Analyze the following document to identify achievements and accomplishments for a military award recommendation. Think through your analysis step-by-step.

First, identify the context: What type of document is this? What time period does it cover? What role/position is described?

Then extract and categorize:
1. **Key Achievements**: Specific accomplishments with measurable outcomes
2. **Leadership Impact**: Examples of leading teams, mentoring, or organizational influence  
3. **Innovation & Initiative**: Process improvements, new solutions, or creative problem-solving
4. **Operational Excellence**: Mission success, readiness improvements, or performance metrics
5. **Challenges Overcome**: Difficult situations handled effectively
6. **Scope of Impact**: Unit-level, command-level, service-wide, or joint/international

For each achievement found:
- Explain WHY it's significant
- Identify the IMPACT (quantify when possible)
- Note any RECOGNITION already received

Document content:
{text}

Provide a comprehensive analysis that helps build a strong award recommendation. If information is missing, suggest specific questions to ask."""

            messages = [
                {"role": "system", "content": "You are a military achievement analyst helping to identify accomplishments for award recommendations. Be specific and focus on concrete achievements."},
                {"role": "user", "content": analysis_prompt.format(text=text[:5000])}  # Limit text to avoid token limits
            ]
            
            response = client.chat_completion(messages)
            return response.get("content", "No achievements identified.")
            
        except Exception as e:
            logger.error(f"Error analyzing document: {e}")
            return f"Error analyzing document: {str(e)}"


# Global document processor instance
document_processor = DocumentProcessor()