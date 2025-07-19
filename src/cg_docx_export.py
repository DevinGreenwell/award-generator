"""
Coast Guard Compliant DOCX Export Module
Generates properly formatted award citations and packages.
"""

import os
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from typing import Dict, Optional


def generate_cg_compliant_docx(export_data: Dict) -> bytes:
    """
    Generate a Coast Guard compliant DOCX file for the award package.
    
    Args:
        export_data: Dictionary containing all award data
        
    Returns:
        Bytes of the DOCX file
    """
    doc = Document()
    
    # Get award data
    finalized = export_data.get('finalized_award')
    recommendation = export_data.get('recommendation')
    awardee_info = export_data.get('awardee_info', {})
    
    # Determine award type
    if finalized:
        award_type = finalized.get('award', '')
        citation_text = finalized.get('citation', '')
    elif recommendation:
        award_type = recommendation.get('award', '')
        citation_text = None  # Will generate if not finalized
    else:
        raise ValueError("No award data found in export")
    
    # Check if this is above MSM level
    above_msm = award_type in ['Distinguished Service Medal', 'Legion of Merit']
    
    # Configure document for landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    
    # Set margins according to CG standards
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(2)  # Space for seal
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Add citation page
    _add_citation_page(doc, award_type, awardee_info, citation_text, export_data)
    
    # Add a new section for the award package (portrait)
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)
    new_section.left_margin = Inches(1.25)
    new_section.right_margin = Inches(1.25)
    
    # Add award package cover page
    _add_package_cover(doc, award_type, awardee_info)
    
    # Add Summary of Action if required
    if _requires_summary_of_action(award_type):
        doc.add_page_break()
        _add_summary_of_action(doc, award_type, export_data)
    
    # Save to bytes
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()


def _add_citation_page(doc: Document, award_type: str, awardee_info: Dict, 
                       citation_text: Optional[str], export_data: Dict):
    """Add the citation page with proper formatting."""
    
    # Add header with award type
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Format the header based on award type
    header_text = f"CITATION TO ACCOMPANY THE AWARD OF"
    header_run = header.add_run(header_text)
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(12)
    
    # Add award name line
    award_line = doc.add_paragraph()
    award_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Format award name (add "THE" prefix for most awards)
    if award_type.startswith("Coast Guard"):
        award_name = f"THE {award_type.upper()}"
    elif award_type in ["Legion of Merit", "Bronze Star Medal"]:
        award_name = f"THE {award_type.upper()}"
    else:
        award_name = award_type.upper()
    
    award_run = award_line.add_run(award_name)
    award_run.font.name = 'Times New Roman'
    award_run.font.size = Pt(12)
    award_run.font.bold = True
    
    # Add "TO" line
    to_line = doc.add_paragraph()
    to_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    to_run = to_line.add_run("TO")
    to_run.font.name = 'Times New Roman'
    to_run.font.size = Pt(12)
    
    # Add blank line
    doc.add_paragraph()
    
    # Add member name and rank
    name_line = doc.add_paragraph()
    name_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name = awardee_info.get('name', '').upper()
    rank = awardee_info.get('rank', '').upper()
    
    if rank:
        name_text = f"{name}\n{rank}"
    else:
        name_text = name
        
    name_run = name_line.add_run(name_text)
    name_run.font.name = 'Times New Roman'
    name_run.font.size = Pt(12)
    name_run.font.bold = True
    
    # Add "UNITED STATES COAST GUARD" line
    uscg_line = doc.add_paragraph()
    uscg_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uscg_run = uscg_line.add_run("UNITED STATES COAST GUARD")
    uscg_run.font.name = 'Times New Roman'
    uscg_run.font.size = Pt(12)
    uscg_run.font.bold = True
    
    # Add blank line before citation
    doc.add_paragraph()
    
    # If no citation text provided, generate one
    if not citation_text:
        from citation_formatter import CitationFormatter
        formatter = CitationFormatter()
        achievement_data = export_data.get('achievement_data', {})
        citation_text = formatter.format_citation(award_type, awardee_info, achievement_data)
    
    # Add the citation text
    citation_para = doc.add_paragraph()
    citation_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Format the citation with proper font
    run = citation_para.add_run(citation_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    
    # Add Operational Distinguishing Device line if applicable
    if _requires_operational_device(award_type):
        doc.add_paragraph()  # Blank line
        device_para = doc.add_paragraph()
        device_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        device_run = device_para.add_run("The Operational Distinguishing Device is authorized.")
        device_run.font.name = 'Times New Roman'
        device_run.font.size = Pt(11)
        device_run.font.bold = True
    
    # Note: In actual implementation, the seal would be added as an image
    # positioned in the lower left corner


def _add_package_cover(doc: Document, award_type: str, awardee_info: Dict):
    """Add the award package cover page."""
    
    # Header
    header = doc.add_heading('UNITED STATES COAST GUARD', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Award type
    award_heading = doc.add_heading(f'{award_type.upper()} RECOMMENDATION', 1)
    award_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    doc.add_paragraph()
    
    # Awardee information
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name = awardee_info.get('name', '')
    rank = awardee_info.get('rank', '')
    
    if rank:
        info_para.add_run(f'{rank} {name}\n').bold = True
    else:
        info_para.add_run(f'{name}\n').bold = True
        
    unit = awardee_info.get('unit', '')
    if unit:
        info_para.add_run(f'{unit}\n')
        
    # Date
    doc.add_paragraph()
    date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_summary_of_action(doc: Document, award_type: str, export_data: Dict):
    """Add Summary of Action page for awards that require it."""
    
    # Header
    header = doc.add_heading('SUMMARY OF ACTION', 1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Award and awardee info
    awardee_info = export_data.get('awardee_info', {})
    achievement_data = export_data.get('achievement_data', {})
    
    info_para = doc.add_paragraph()
    info_para.add_run('AWARD: ').bold = True
    info_para.add_run(award_type)
    
    info_para = doc.add_paragraph()
    info_para.add_run('NAME: ').bold = True
    name = awardee_info.get('name', '')
    rank = awardee_info.get('rank', '')
    info_para.add_run(f'{rank} {name}' if rank else name)
    
    info_para = doc.add_paragraph()
    info_para.add_run('UNIT: ').bold = True
    info_para.add_run(awardee_info.get('unit', ''))
    
    info_para = doc.add_paragraph()
    info_para.add_run('PERIOD: ').bold = True
    info_para.add_run(achievement_data.get('time_period', ''))
    
    doc.add_paragraph()
    
    # Summary section
    summary_heading = doc.add_heading('SUMMARY', 2)
    
    # Build summary from achievements and impacts
    achievements = achievement_data.get('achievements', [])
    impacts = achievement_data.get('impacts', [])
    
    if achievements:
        doc.add_paragraph('KEY ACHIEVEMENTS:')
        for achievement in achievements[:5]:
            bullet = doc.add_paragraph(f'• {achievement}')
            bullet.paragraph_format.left_indent = Inches(0.5)
    
    if impacts:
        doc.add_paragraph()
        doc.add_paragraph('MEASURABLE IMPACTS:')
        for impact in impacts[:5]:
            bullet = doc.add_paragraph(f'• {impact}')
            bullet.paragraph_format.left_indent = Inches(0.5)
    
    # Justification
    if achievement_data.get('justification'):
        doc.add_paragraph()
        just_heading = doc.add_heading('JUSTIFICATION', 2)
        doc.add_paragraph(achievement_data.get('justification'))


def _requires_summary_of_action(award_type: str) -> bool:
    """Check if the award type requires a Summary of Action."""
    return award_type in [
        'Distinguished Service Medal',
        'Legion of Merit', 
        'Meritorious Service Medal',
        'Coast Guard Unit Commendation'
    ]


def _requires_operational_device(award_type: str) -> bool:
    """Check if the award type may include the Operational Distinguishing Device."""
    # Based on the examples, these awards sometimes include the device
    return award_type in [
        'Coast Guard Commendation Medal',
        'Coast Guard Achievement Medal',
        'Coast Guard Unit Commendation',
        'Coast Guard Meritorious Unit Commendation',
        'Coast Guard Meritorious Team Commendation'
    ]


def set_cg_document_styles(doc: Document):
    """Set document styles to match CG requirements."""
    styles = doc.styles
    
    # Normal style
    try:
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
    except:
        pass
    
    # Heading styles
    try:
        for i in range(1, 4):
            heading = styles[f'Heading {i}']
            heading.font.name = 'Times New Roman'
            heading.font.bold = True
            if i == 1:
                heading.font.size = Pt(14)
            elif i == 2:
                heading.font.size = Pt(13)
            else:
                heading.font.size = Pt(12)
    except:
        pass