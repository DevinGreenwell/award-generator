"""
Main Flask application for the Coast Guard Award Generator.
"""

import os
import sys
import json
import logging
from datetime import datetime
from io import BytesIO
from functools import wraps
from pathlib import Path

# Add the src directory to Python path for imports to work in both local and deployed environments
# This handles cases where the app is run from different directories
current_dir = Path(__file__).parent
if current_dir.name == 'src':
    # We're in the src directory, add parent to path
    sys.path.insert(0, str(current_dir.parent))
    # Also add src itself to path
    sys.path.insert(0, str(current_dir))
else:
    # We might be running from project root or elsewhere
    src_dir = current_dir / 'src'
    if src_dir.exists():
        sys.path.insert(0, str(src_dir))

from flask import Flask, render_template, request, jsonify, session, send_file, make_response
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Try different import strategies for compatibility
try:
    # Try direct import first (works in most cases)
    from award_engine import AwardEngine, AwardEngineError, InsufficientDataError
    from openai_client import OpenAIClient
    from config import current_config, setup_logging
    from validation import (
        ValidationError, AwardeeInfoValidator, AchievementDataValidator,
        MessageValidator, ExportRequestValidator, SessionDataValidator
    )
    from session_manager import (
        store_session_data, get_session_data, clear_session_data,
        get_or_create_session_id
    )
except ImportError:
    # If direct import fails, try with src prefix
    try:
        from src.award_engine import AwardEngine, AwardEngineError, InsufficientDataError
        from src.openai_client import OpenAIClient
        from src.config import current_config, setup_logging
        from src.validation import (
            ValidationError, AwardeeInfoValidator, AchievementDataValidator,
            MessageValidator, ExportRequestValidator, SessionDataValidator
        )
        from src.session_manager import (
            store_session_data, get_session_data, clear_session_data,
            get_or_create_session_id
        )
    except ImportError as e:
        # Log the error and re-raise with helpful message
        import traceback
        print(f"Import error: {e}")
        print(f"Current directory: {os.getcwd()}")
        print(f"Python path: {sys.path}")
        print(f"Directory contents: {os.listdir('.')}")
        traceback.print_exc()
        raise ImportError(
            f"Failed to import required modules. "
            f"Current directory: {os.getcwd()}, "
            f"__file__ location: {__file__}"
        )

# Set up logging
logger = setup_logging()

# Validate configuration
try:
    current_config.validate()
except ValueError as e:
    logger.error(f"Configuration error: {e}")
    raise

# Configure Flask app
app = Flask(__name__, 
           static_folder='static',
           static_url_path='/static',
           template_folder='templates')

# Apply configuration
app.config.from_object(current_config)

# Set up CORS
CORS(app, origins=current_config.CORS_ORIGINS)

# Initialize services
try:
    award_engine = AwardEngine()
    openai_client = OpenAIClient()
    logger.info("Services initialized successfully")
    
    # Clean up old sessions on startup
    from session_manager import session_manager
    session_manager.cleanup_old_sessions()
except Exception as e:
    logger.error(f"Failed to initialize services: {e}")
    raise


def handle_errors(f):
    """Decorator to handle errors consistently across endpoints."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        try:
            return f(*args, **kwargs)
        except ValidationError as e:
            logger.warning(f"Validation error in {f.__name__}: {e}")
            return jsonify({"error": str(e), "type": "validation"}), 400
        except AwardEngineError as e:
            logger.error(f"Award engine error in {f.__name__}: {e}")
            return jsonify({"error": str(e), "type": "award_engine"}), 500
        except Exception as e:
            logger.error(f"Unexpected error in {f.__name__}: {e}", exc_info=True)
            return jsonify({"error": "An unexpected error occurred", "type": "internal"}), 500
    return decorated_function


@app.route('/')
def index():
    """Serve the main application page."""
    return render_template('index.html')


@app.route('/api/session/clear', methods=['POST'])
@handle_errors
def clear_session():
    """Clear the current session data."""
    # Clear file-based session data
    clear_session_data(session)
    
    # Clear any remaining cookie session data
    session.clear()
    
    logger.info("Session cleared")
    return jsonify({
        'success': True,
        'message': 'Session cleared successfully'
    })


@app.route('/api/chat', methods=['POST'])
@handle_errors
def api_chat():
    """Handle chat messages and store conversation history."""
    # Validate input
    data = MessageValidator.validate(request.get_json())
    message = data['message']
    
    # Get existing messages from file-based session
    messages = get_session_data(session, 'messages') or []
    
    # Add the new user message
    messages.append({
        "role": "user", 
        "content": message,
        "timestamp": datetime.now().isoformat()
    })
    
    # Check if we have document context
    document_text = get_session_data(session, 'document_text')
    document_analysis = get_session_data(session, 'document_analysis')
    
    # Prepare messages for OpenAI
    system_content = "You are a helpful assistant helping to document Coast Guard achievements for award recommendations. Acknowledge the user's input and encourage them to continue sharing details."
    
    # Add document context to system message if available
    if document_text:
        system_content += f"\n\nIMPORTANT: You have access to a previously uploaded document. When the user asks about specific details from 'the document' or 'the uploaded document', you can refer to this content:\n\n{document_text[:5000]}"
        if len(document_text) > 5000:
            system_content += f"\n[Document continues - {len(document_text)} total characters]"
        system_content += "\n\nWhen answering questions about the document, cite specific sections or details from the above content."
    
    openai_messages = [
        {"role": "system", "content": system_content}
    ]
    
    # Add conversation history
    for msg in messages:
        if msg.get('role') in ['user', 'assistant'] and msg.get('content'):
            openai_messages.append({
                "role": msg['role'],
                "content": msg['content']
            })
    
    logger.debug(f"Sending {len(openai_messages)} messages to OpenAI")
    
    # Generate AI response
    ai_response = openai_client.chat_completion(openai_messages)
    
    # Add AI response to messages
    messages.append({
        "role": "assistant",
        "content": ai_response.get("content", "I understand. Please continue."),
        "timestamp": datetime.now().isoformat()
    })
    
    # Store updated messages in file-based session
    store_session_data(session, 'messages', messages)
    
    logger.info(f"Chat interaction completed. Total messages: {len(messages)}")
    
    return jsonify({
        "success": True,
        "message": ai_response.get("content", "I understand. Please continue."),
        "response": ai_response.get("content", "I understand. Please continue."),  # Keep for backwards compatibility
        "message_count": len(messages)
    })


@app.route('/api/recommend', methods=['POST'])
@handle_errors
def api_recommend():
    """Generate award recommendation based on conversation."""
    data = request.get_json()
    
    # Validate awardee info if provided
    awardee_info = {}
    if 'awardee_info' in data:
        awardee_info = AwardeeInfoValidator.validate(data['awardee_info'])
    
    # Get ALL messages from the file-based session
    messages = get_session_data(session, 'messages') or []
    
    if not messages:
        raise ValidationError("No achievements have been added yet. Please describe some achievements using the chat interface first.")
    
    # Filter to only user messages for verification
    user_messages = [msg for msg in messages if msg.get('role') == 'user']
    
    if not user_messages:
        raise ValidationError("No achievement descriptions found. Please describe your accomplishments using the chat interface.")
    
    logger.info(f"Analyzing {len(user_messages)} user messages with {len(messages)} total messages")
    
    # Analyze messages
    achievement_data = openai_client.analyze_achievements(messages, awardee_info)
    
    # Validate achievement data
    achievement_data = AchievementDataValidator.validate(achievement_data)
    
    # Store the analysis in file-based session
    store_session_data(session, 'achievement_data', achievement_data)
    store_session_data(session, 'awardee_info', awardee_info)
    
    # Score the achievements
    scores = award_engine.score_achievements(achievement_data)
    
    # Get award recommendation
    recommendation = award_engine.recommend_award(scores)
    award = recommendation["award"]
    
    # Generate explanation
    explanation = award_engine.generate_explanation(award, achievement_data, scores)
    
    # Generate improvement suggestions
    suggestions = award_engine.generate_improvement_suggestions(award, achievement_data)
    
    # Store recommendation in file-based session
    recommendation_data = {
        'award': award,
        'explanation': explanation,
        'achievement_data': achievement_data,
        'scores': scores,
        'suggestions': suggestions
    }
    store_session_data(session, 'recommendation', recommendation_data)
    
    logger.info(f"Generated recommendation: {award} with score {recommendation['score']}")
    
    return jsonify({
        "success": True,
        "award": award,
        "explanation": explanation,
        "scores": scores,
        "achievement_data": achievement_data,
        "suggestions": suggestions,
        "message_count": len(messages)
    })


@app.route('/api/refresh', methods=['POST'])
@handle_errors
def api_refresh():
    """Refresh the award recommendation with alternative analysis."""
    data = request.get_json()
    
    # Validate awardee info if provided
    awardee_info = get_session_data(session, 'awardee_info') or {}
    if 'awardee_info' in data:
        awardee_info = AwardeeInfoValidator.validate(data['awardee_info'])
    
    # Get messages from file-based session
    messages = get_session_data(session, 'messages') or []
    
    if not messages:
        raise ValidationError("No conversation found. Please start a new conversation.")
    
    logger.info("Refreshing recommendation with alternative analysis")
    
    # Analyze messages with refresh flag
    achievement_data = openai_client.analyze_achievements(messages, awardee_info, refresh=True)
    
    # Validate achievement data
    achievement_data = AchievementDataValidator.validate(achievement_data)
    
    # Update file-based session
    store_session_data(session, 'achievement_data', achievement_data)
    
    # Score and recommend
    scores = award_engine.score_achievements(achievement_data)
    recommendation = award_engine.recommend_award(scores)
    award = recommendation["award"]
    
    # Generate explanation and suggestions
    explanation = award_engine.generate_explanation(award, achievement_data, scores)
    suggestions = award_engine.generate_improvement_suggestions(award, achievement_data)
    
    # Update recommendation in file-based session
    recommendation_data = {
        'award': award,
        'explanation': explanation,
        'achievement_data': achievement_data,
        'scores': scores,
        'suggestions': suggestions
    }
    store_session_data(session, 'recommendation', recommendation_data)
    
    logger.info(f"Refreshed recommendation: {award}")
    
    return jsonify({
        "success": True,
        "award": award,
        "explanation": explanation,
        "scores": scores,
        "achievement_data": achievement_data,
        "suggestions": suggestions
    })


@app.route('/api/improve', methods=['POST'])
@handle_errors
def api_improve():
    """Generate improvement suggestions for the current award recommendation."""
    data = request.get_json()
    current_award = data.get('current_award', '')
    
    # Get current award from session if not provided
    if not current_award:
        recommendation = session.get('recommendation', {})
        current_award = recommendation.get('award', '')
    
    if not current_award:
        raise ValidationError('No current award found. Please generate a recommendation first.')
    
    # Get current session data from file-based session
    achievement_data = get_session_data(session, 'achievement_data') or {}
    awardee_info = get_session_data(session, 'awardee_info') or {}
    
    if not achievement_data:
        raise ValidationError('No achievement data found. Please generate a recommendation first.')
    
    logger.info(f"Generating improvement suggestions for {current_award}")
    
    # Generate improvement suggestions
    suggestions = openai_client.generate_improvement_suggestions(
        current_award, achievement_data, awardee_info
    )
    
    # Get current scores for reference
    current_scores = award_engine.score_achievements(achievement_data)
    
    # Store improvement suggestions in file-based session
    improvement_data = {
        'current_award': current_award,
        'suggestions': suggestions,
        'current_scores': current_scores
    }
    store_session_data(session, 'improvement_suggestions', improvement_data)
    
    return jsonify({
        'success': True,
        'current_award': current_award,
        'suggestions': suggestions,
        'current_scores': current_scores
    })


@app.route('/api/finalize', methods=['POST'])
@handle_errors
def api_finalize():
    """Generate final award citation."""
    data = request.get_json()
    award = data.get('award')
    
    # Get award from session if not provided
    if not award:
        recommendation = session.get('recommendation', {})
        award = recommendation.get('award', '')
    
    if not award:
        raise ValidationError('No award specified. Please generate a recommendation first.')
    
    # Get current achievement data from file-based session
    achievement_data = get_session_data(session, 'achievement_data') or {}
    awardee_info = get_session_data(session, 'awardee_info') or {}
    
    if not achievement_data:
        raise ValidationError('No achievement data found. Please generate a recommendation first.')
    
    logger.info(f"Generating citation for award: {award}")
    
    # Generate final citation
    citation = openai_client.draft_award(award, achievement_data, awardee_info)
    
    # Store finalized award in file-based session
    finalized_data = {
        'award': award,
        'citation': citation,
        'finalized_at': datetime.now().isoformat()
    }
    store_session_data(session, 'finalized_award', finalized_data)
    
    return jsonify({
        'success': True,
        'award': award,
        'citation': citation
    })


@app.route('/api/upload', methods=['POST'])
@handle_errors
def api_upload():
    """Handle document upload and text extraction."""
    # Check if file is in request
    if 'file' not in request.files:
        raise ValidationError('No file provided')
    
    file = request.files['file']
    
    # Process the file
    from document_processor import document_processor
    success, message, extracted_text = document_processor.process_file(file)
    
    if success:
        logger.info(f"Successfully processed file: {file.filename}")
        
        # Analyze the document for achievements
        analysis = document_processor.analyze_document_for_achievements(extracted_text)
        
        # Store document analysis in session for later use
        if analysis:
            store_session_data(session, 'document_analysis', analysis)
            # Also store the original extracted text for retrieval
            store_session_data(session, 'document_text', extracted_text)
            
            # Add a user message with the document analysis
            messages = get_session_data(session, 'messages') or []
            messages.append({
                "role": "user",
                "content": f"Document content and analysis: {analysis}",
                "timestamp": datetime.now().isoformat()
            })
            store_session_data(session, 'messages', messages)
        
        return jsonify({
            'success': True,
            'message': f"Successfully analyzed {file.filename}",
            'analysis': analysis,
            'extracted_text': analysis  # Send analysis as extracted_text for compatibility
        })
    else:
        logger.warning(f"Failed to process file: {message}")
        return jsonify({
            'success': False,
            'error': message
        }), 400


@app.route('/api/export', methods=['POST'])
@handle_errors
def api_export():
    """Generate export data in multiple formats including DOCX."""
    # Validate request
    data = ExportRequestValidator.validate(request.get_json())
    export_format = data['format']
    
    # Update awardee info if provided
    if data['awardee_info']:
        store_session_data(session, 'awardee_info', data['awardee_info'])
    
    # Gather all session data from file-based session
    finalized = get_session_data(session, 'finalized_award')
    recommendation = get_session_data(session, 'recommendation')
    achievement_data = get_session_data(session, 'achievement_data') or {}
    awardee_info = get_session_data(session, 'awardee_info') or {}
    messages = get_session_data(session, 'messages') or []
    
    # Prepare export data
    export_data = {
        'generated_at': datetime.now().isoformat(),
        'awardee_info': awardee_info,
        'messages': messages,
        'achievement_data': achievement_data,
        'recommendation': recommendation,
        'finalized_award': finalized
    }
    
    # Generate filename
    name = awardee_info.get('name', 'Unknown').replace(' ', '_')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    logger.info(f"Exporting award package in {export_format} format")
    
    if export_format == 'docx':
        filename = f"award_package_{name}_{timestamp}.docx"
        # Store export data in file-based session for download endpoint
        store_session_data(session, 'export_data', export_data)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'download_url': '/api/export/download/docx',
            'format': 'docx'
        })
        
    elif export_format == 'json':
        filename = f"award_package_{name}_{timestamp}.json"
        content = json.dumps(export_data, indent=2, default=str)
        mimetype = 'application/json'
        
    elif export_format == 'txt':
        filename = f"award_package_{name}_{timestamp}.txt"
        content = generate_text_export(export_data)
        mimetype = 'text/plain'
        
    else:
        raise ValidationError(f'Unsupported export format: {export_format}')
    
    # For non-DOCX formats, return content directly
    return jsonify({
        'success': True,
        'filename': filename,
        'content': content,
        'mimetype': mimetype,
        'size': len(content.encode('utf-8'))
    })


@app.route('/api/export/download/docx', methods=['GET'])
@handle_errors
def download_docx():
    """Download the generated DOCX file."""
    # Get export data from file-based session
    export_data = get_session_data(session, 'export_data')
    if not export_data:
        raise ValidationError("No export data found. Please generate an export first.")
    
    # Generate DOCX
    doc_bytes = generate_docx_export(export_data)
    
    # Create filename
    awardee_info = export_data.get('awardee_info', {})
    name = awardee_info.get('name', 'Unknown').replace(' ', '_')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"award_package_{name}_{timestamp}.docx"
    
    # Create response
    response = make_response(doc_bytes)
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
    response.headers['Content-Length'] = len(doc_bytes)
    
    logger.info(f"DOCX download completed: {filename}")
    
    return response


@app.route('/api/session', methods=['GET'])
@handle_errors
def get_session():
    """Get current session data."""
    session_data = {
        "session_id": get_session_data(session, 'session_id') or 'default',
        "session_name": get_session_data(session, 'session_name') or '',
        "messages": get_session_data(session, 'messages') or [],
        "awardee_info": get_session_data(session, 'awardee_info') or {},
        "recommendation": get_session_data(session, 'recommendation'),
        "finalized_award": get_session_data(session, 'finalized_award')
    }
    return jsonify(session_data)


@app.route('/api/session', methods=['POST'])
@handle_errors
def save_session():
    """Save session data."""
    # Validate input
    data = SessionDataValidator.validate(request.get_json())
    
    # Save relevant fields to file-based session
    for key in ['session_id', 'session_name', 'messages', 'awardee_info']:
        if key in data:
            store_session_data(session, key, data[key])
    
    session.permanent = True
    
    logger.info(f"Session saved: {data.get('session_name', 'Unnamed')}")
    
    return jsonify({
        "status": "success", 
        "message": "Session data saved.", 
        "session_id": session.get('session_id', 'default')
    })


@app.route('/api/debug/session', methods=['GET'])
def debug_session():
    """Debug endpoint to check session state (development only)."""
    if not app.debug:
        return jsonify({"error": "Debug endpoint not available in production"}), 403
    
    # Get session data
    session_data = get_session_data(session) or {}
    messages = get_session_data(session, 'messages') or []
    
    return jsonify({
        "session_keys": list(session_data.keys()),
        "messages_count": len(messages),
        "has_achievement_data": 'achievement_data' in session_data,
        "has_recommendation": 'recommendation' in session_data,
        "session_id": get_session_data(session, 'session_id') or 'No ID',
        "recent_messages": [
            {
                "role": msg.get('role'),
                "content": msg.get('content', '')[:100] + "..." if len(msg.get('content', '')) > 100 else msg.get('content', ''),
                "timestamp": msg.get('timestamp')
            }
            for msg in messages[-5:]  # Last 5 messages
        ]
    })


def generate_docx_export(export_data):
    """Generate a professionally formatted DOCX file."""
    doc = Document()
    
    # Set up document margins and styles
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Add custom styles
    add_custom_styles(doc)
    
    # Get data
    awardee_info = export_data.get('awardee_info', {})
    finalized = export_data.get('finalized_award')
    recommendation = export_data.get('recommendation')
    achievement_data = export_data.get('achievement_data', {})
    
    # Document header
    header = doc.add_heading('UNITED STATES COAST GUARD', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subheader = doc.add_heading('Award Recommendation Package', 1)
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Awardee Information Section
    if awardee_info:
        doc.add_heading('I. AWARDEE INFORMATION', 2)
        
        # Create a formatted table for awardee info
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        info_fields = [
            ('Name', awardee_info.get('name', '')),
            ('Rank/Rate', awardee_info.get('rank', '')),
            ('Unit', awardee_info.get('unit', '')),
            ('Position', awardee_info.get('position', '')),
            ('EMPLID', awardee_info.get('service_number', '')),
        ]
        
        for field, value in info_fields:
            if value:
                row = table.add_row()
                row.cells[0].text = field + ':'
                row.cells[1].text = str(value)
                # Bold the field names
                row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
    
    # Award Recommendation Section
    award_name = ''
    if finalized:
        award_name = finalized.get('award', '')
        doc.add_heading('II. FINAL AWARD RECOMMENDATION', 2)
        
        award_para = doc.add_paragraph()
        award_para.add_run('Recommended Award: ').bold = True
        award_para.add_run(award_name)
        
    elif recommendation:
        award_name = recommendation.get('award', '')
        doc.add_heading('II. AWARD RECOMMENDATION', 2)
        
        award_para = doc.add_paragraph()
        award_para.add_run('Recommended Award: ').bold = True
        award_para.add_run(award_name)
    
    doc.add_paragraph()
    
    # Citation Section
    if finalized and finalized.get('citation'):
        doc.add_heading('III. CITATION', 2)
        citation_para = doc.add_paragraph(finalized.get('citation'))
        citation_para.style = 'Citation'
        doc.add_paragraph()
    
    # Achievement Summary
    if achievement_data:
        doc.add_heading('IV. ACHIEVEMENT SUMMARY', 2)
        
        # Key Achievements
        achievements = achievement_data.get('achievements', [])
        if achievements:
            doc.add_heading('A. Key Achievements', 3)
            for i, achievement in enumerate(achievements[:8], 1):
                para = doc.add_paragraph(f"{i}. {achievement}")
                para.style = 'List Number'
            doc.add_paragraph()
        
        # Measurable Impact
        impacts = achievement_data.get('impacts', [])
        if impacts:
            doc.add_heading('B. Measurable Impact', 3)
            for impact in impacts[:6]:
                para = doc.add_paragraph(impact, style='List Bullet')
            doc.add_paragraph()
        
        # Leadership Details
        leadership = achievement_data.get('leadership_details', [])
        if leadership:
            doc.add_heading('C. Leadership Demonstrated', 3)
            for detail in leadership[:5]:
                para = doc.add_paragraph(detail, style='List Bullet')
            doc.add_paragraph()
        
        # Innovation and Initiative
        innovations = achievement_data.get('innovation_details', [])
        if innovations:
            doc.add_heading('D. Innovation and Initiative', 3)
            for innovation in innovations[:5]:
                para = doc.add_paragraph(innovation, style='List Bullet')
            doc.add_paragraph()
        
        # Challenges Overcome
        challenges = achievement_data.get('challenges', [])
        if challenges:
            doc.add_heading('E. Challenges Overcome', 3)
            for challenge in challenges[:5]:
                para = doc.add_paragraph(challenge, style='List Bullet')
            doc.add_paragraph()
        
        # Scope and Time Period
        scope = achievement_data.get('scope', '')
        time_period = achievement_data.get('time_period', '')
        
        if scope and scope != 'Not specified':
            para = doc.add_paragraph()
            para.add_run('Scope of Impact: ').bold = True
            para.add_run(scope)
        
        if time_period and time_period != 'Not specified':
            para = doc.add_paragraph()
            para.add_run('Time Period: ').bold = True
            para.add_run(time_period)
        
        doc.add_paragraph()
    
    # Scoring Analysis
    if recommendation and recommendation.get('scores'):
        doc.add_heading('V. SCORING ANALYSIS', 2)
        
        scores = recommendation.get('scores', {})
        total_score = scores.get('total_weighted', 0)
        
        # Overall score
        score_para = doc.add_paragraph()
        score_para.add_run('Overall Score: ').bold = True
        score_para.add_run(f"{total_score}/100")
        
        # Detailed scoring table
        doc.add_paragraph('Detailed Scoring Breakdown:')
        
        score_table = doc.add_table(rows=1, cols=2)
        score_table.style = 'Table Grid'
        
        # Header row
        header_row = score_table.rows[0]
        header_row.cells[0].text = 'Criterion'
        header_row.cells[1].text = 'Score (out of 5)'
        
        # Make header bold
        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Add score rows
        for criterion, score in scores.items():
            if criterion != 'total_weighted' and score > 0:
                row = score_table.add_row()
                row.cells[0].text = criterion.replace('_', ' ').title()
                row.cells[1].text = f"{score}/5.0"
        
        doc.add_paragraph()
    
    # Justification
    if achievement_data.get('justification'):
        doc.add_heading('VI. JUSTIFICATION', 2)
        justification_para = doc.add_paragraph(achievement_data.get('justification'))
        doc.add_paragraph()
    
    # Document footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run('Created by Coast Guard Award Generator on ').italic = True
    footer_para.add_run(datetime.now().strftime('%B %d, %Y at %I:%M %p')).italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()


def add_custom_styles(doc):
    """Add custom styles to the document."""
    styles = doc.styles
    
    # Citation style
    try:
        citation_style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
        citation_style.font.size = Pt(11)
        citation_style.font.italic = True
        citation_style.paragraph_format.left_indent = Inches(0.5)
        citation_style.paragraph_format.right_indent = Inches(0.5)
        citation_style.paragraph_format.space_before = Pt(6)
        citation_style.paragraph_format.space_after = Pt(6)
    except:
        pass  # Style might already exist
    
    # Adjust existing styles
    try:
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        
        heading1_style = styles['Heading 1']
        heading1_style.font.name = 'Times New Roman'
        heading1_style.font.size = Pt(14)
        heading1_style.font.bold = True
        
        heading2_style = styles['Heading 2']
        heading2_style.font.name = 'Times New Roman'
        heading2_style.font.size = Pt(13)
        heading2_style.font.bold = True
        
        heading3_style = styles['Heading 3']
        heading3_style.font.name = 'Times New Roman'
        heading3_style.font.size = Pt(12)
        heading3_style.font.bold = True
    except:
        pass  # Styles might not be available


def generate_text_export(export_data):
    """Generate a human-readable text export."""
    lines = []
    
    # Header
    lines.append("=" * 60)
    lines.append("COAST GUARD AWARD PACKAGE EXPORT")
    lines.append("=" * 60)
    lines.append(f"Generated: {export_data.get('generated_at', 'Unknown')}")
    lines.append("")
    
    # Awardee Information
    awardee_info = export_data.get('awardee_info', {})
    if awardee_info:
        lines.append("AWARDEE INFORMATION:")
        lines.append("-" * 30)
        for key, value in awardee_info.items():
            if value:
                lines.append(f"{key.replace('_', ' ').title()}: {value}")
        lines.append("")
    
    # Final Award
    finalized = export_data.get('finalized_award')
    if finalized:
        lines.append("FINAL AWARD:")
        lines.append("-" * 30)
        lines.append(f"Award: {finalized.get('award', 'N/A')}")
        lines.append("")
        lines.append("CITATION:")
        lines.append(finalized.get('citation', 'No citation available'))
        lines.append("")
    
    # Achievement Data
    achievement_data = export_data.get('achievement_data', {})
    if achievement_data:
        lines.append("ACHIEVEMENT ANALYSIS:")
        lines.append("-" * 30)
        
        # Key sections
        sections = [
            ('achievements', 'Key Achievements'),
            ('impacts', 'Measurable Impacts'),
            ('leadership_details', 'Leadership Details'),
            ('innovation_details', 'Innovation & Initiative'),
            ('challenges', 'Challenges Overcome')
        ]
        
        for field, title in sections:
            items = achievement_data.get(field, [])
            if items:
                lines.append(f"{title}:")
                for item in items[:5]:  # Limit to top 5
                    lines.append(f"  • {item}")
                lines.append("")
    
    return "\n".join(lines)


@app.errorhandler(404)
def not_found(e):
    """Handle 404 errors."""
    return jsonify({"error": "Endpoint not found"}), 404


@app.errorhandler(500)
def server_error(e):
    """Handle 500 errors."""
    logger.error(f"Server error: {e}")
    return jsonify({"error": "Internal server error"}), 500


if __name__ == "__main__":
    # For local development
    app.run(debug=True, host='0.0.0.0', port=5000)
else:
    # For production (Railway will use this)
    app.config['DEBUG'] = False